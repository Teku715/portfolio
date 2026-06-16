from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent
PHOTO_PATH = BASE_DIR / 'photo.jpg'
OUTPUT_PATH = BASE_DIR / 'resume_final.docx'

doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


def set_run_font(run, size=10, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color


def add_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=13, bold=True, color=RGBColor(0x1a, 0x1a, 0x2e))
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2563EB')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_item(doc, title, meta='', content=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    if meta:
        run1 = p.add_run(title)
        set_run_font(run1, size=11, bold=True)
        run2 = p.add_run(f'  {meta}')
        set_run_font(run2, size=10, color=RGBColor(0x88, 0x88, 0x88))
    else:
        run = p.add_run(title)
        set_run_font(run, size=11, bold=True)
    if content:
        p2 = doc.add_paragraph(content)
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(2)
        for run in p2.runs:
            set_run_font(run, color=RGBColor(0x44, 0x44, 0x44))


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        set_run_font(run, color=RGBColor(0x44, 0x44, 0x44))


table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
tblBorders = OxmlElement('w:tblBorders')
for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    border = OxmlElement(f'w:{border_name}')
    border.set(qn('w:val'), 'none')
    tblBorders.append(border)
tblPr.append(tblBorders)

cell_left = table.cell(0, 0)
cell_right = table.cell(0, 1)

p_name = cell_left.paragraphs[0]
p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p_name.add_run('特列努尔')
set_run_font(run, size=22, bold=True, color=RGBColor(0x1a, 0x1a, 0x2e))

p2 = cell_left.add_paragraph()
run2 = p2.add_run('Java 后端开发实习生  |  南京  |  可远程')
set_run_font(run2, size=12, color=RGBColor(0x25, 0x63, 0xeb))

p3 = cell_left.add_paragraph()
run3 = p3.add_run('18761851809  |  1079086024@qq.com  |  github.com/Teku715')
set_run_font(run3, size=11, color=RGBColor(0x55, 0x55, 0x55))

p4 = cell_left.add_paragraph()
run4 = p4.add_run('南京邮电大学 · 物联网学院 · 网络工程 · 2025级 · 大一')
set_run_font(run4, size=10, color=RGBColor(0x55, 0x55, 0x55))

if PHOTO_PATH.exists():
    p_photo = cell_right.paragraphs[0]
    p_photo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_photo = p_photo.add_run()
    run_photo.add_picture(str(PHOTO_PATH), width=Inches(0.9))

cell_left.width = Inches(5.0)
cell_right.width = Inches(1.3)

doc.add_paragraph()

add_title(doc, '教育背景')
add_item(doc, '南京邮电大学', '2025级 · 大一', '物联网学院 · 网络工程  |  GPA 3.15 / 4.0  |  英语四级已通过')

add_title(doc, '专业技能')
skills = [
    ('Java 后端', 'Spring Boot / RESTful API / MySQL / MyBatis / JWT / Spring AOP'),
    ('AI & 大模型', '硅基流动 Qwen2.5 / SenseVoice / AI 辅助编程 / Copilot / Coze'),
    ('前端 & 工具', 'Vue3 / Element Plus / Vite / Git / Docker / Maven'),
]
for title, desc in skills:
    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run(f'{title}：')
    set_run_font(run1, size=10, bold=True)
    run2 = p.add_run(desc)
    set_run_font(run2, color=RGBColor(0x44, 0x44, 0x44))

add_title(doc, '项目经历')

projects = [
    (
        '七牛云 XEngineer · 丝路声绘 · Silk VoiceBrush（AI 英语口语陪练）',
        '2026.06',
        [
            '独立完成后端架构，集成 SenseVoice 语音识别与 Qwen 大模型，实现多场景口语练习全流程',
            '设计场景对话、按住说话、表达评测、课后总结等核心模块，完成前后端联调与部署',
            '支持面试 / 餐厅 / 会议 / 旅行 / 购物 5 类场景，涵盖 ASR、TTS、角色扮演与打字机回复',
        ],
    ),
    (
        '中兴捧月 · 兴享智家（校园课程助手智能体）',
        '2026.05',
        [
            '独立完成前后端开发，接入 Qwen2.5-7B 实现 AI 对话，支持查课程表、查成绩、智能问答',
            '后端提供 RESTful 接口，前端基于 Vue3 + Element Plus 完成交互与数据展示',
            '完成从需求分析、接口设计到联调上线的完整开发流程',
        ],
    ),
    (
        '星火工作室考核 · CStudyMate C语言 AI 学伴',
        '2026.05',
        [
            '后端基于 Spring Boot + MyBatis 实现 RESTful API，前端完成 AI 问答、练习中心、学情统计等模块',
            '接入 Qwen2.5-7B，设计结构化 Prompt，强制输出概念 / 原理 / 代码 / 易错点',
            '实现问答历史、练习记录等学习闭环功能，完成前后端联调',
        ],
    ),
    (
        'TLIAS 智能学习平台',
        '2026.03',
        [
            '实现基于 JWT 的用户登录认证与令牌鉴权，Filter 拦截未授权请求',
            '使用 Spring AOP 统一记录操作日志，完成部门、员工等管理模块接口开发',
            '独立完成前后端联调与功能测试',
        ],
    ),
    (
        '计算机设计大赛 · 老人信息管理系统',
        '2026.04',
        [
            '独立完成 Java 后端开发，设计并实现 5 个 REST API 接口（增删改查）',
            '使用 MySQL 完成数据持久化，规范接口设计与异常处理',
        ],
    ),
]

for title, meta, bullets in projects:
    add_item(doc, title, meta)
    add_bullets(doc, bullets)

add_title(doc, '个人评价')
evals = [
    '从新疆到南京独立求学，具备从需求分析到部署上线的完整项目经验',
    '熟悉 AI 辅助编程工具（Cursor、Copilot、Coze），能结合大模型 API 快速完成业务开发',
    '学习能力强，大一下学期自学 Spring Boot、Vue3、MyBatis 并落地多个全栈项目',
]
for e in evals:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(e)
    set_run_font(run, color=RGBColor(0x44, 0x44, 0x44))

doc.save(str(OUTPUT_PATH))
print(f'done: {OUTPUT_PATH}')
